import os
from docx import Document as DocxDocument
from langchain_core.documents import Document


def load_docx(file_path: str):
    """Load a DOCX file and return a list of LangChain Document objects."""
    docx_doc = DocxDocument(file_path)

    # Extract paragraphs, skip empty ones
    paragraphs = [p.text.strip() for p in docx_doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    filename = os.path.basename(file_path)

    return [
        Document(
            page_content=full_text,
            metadata={"source": filename, "type": "docx"},
        )
    ]